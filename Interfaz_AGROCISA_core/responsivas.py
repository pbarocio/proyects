import streamlit as st
import pandas as pd
from datetime import datetime
import io
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from database import obtener_conexion

def aplicar_estilos_pantalla():
    st.markdown("""
        <style>
            .appview-container .main .block-container {
                max-width: 100% !important;
                padding-left: 2rem !important;
                padding-right: 2rem !important;
            }
        </style>
    """, unsafe_allow_html=True)

def obtener_empleados_activos_df():
    """Obtiene colaboradores activos con el detalle de sucursal, depto y puesto."""
    try:
        conn = obtener_conexion()
        query = """
            SELECT 
                e.codigo,
                CONCAT(e.nombre, ' ', e.apellido_paterno, ' ', COALESCE(e.apellido_materno, '')) AS nombre_completo,
                s.nombre_sucursal AS sucursal,
                d.nombre_departamento AS departamento,
                p.nombre_puesto AS puesto
            FROM empleados e
            LEFT JOIN sucursales s ON e.id_sucursal = s.id_sucursal
            LEFT JOIN departamentos d ON e.id_departamento = d.id_departamento
            LEFT JOIN puestos p ON e.id_puesto = p.id_puesto
            WHERE e.id_estatus_empleado = 1
            ORDER BY nombre_completo ASC
        """
        df = pd.read_sql(query, conn)
        conn.close()
        df["codigo_str"] = df["codigo"].astype(str).str.strip()
        return df
    except Exception as e:
        st.error(f"⚠️ Error al cargar empleados: {e}")
        return pd.DataFrame()

def obtener_equipos_disponibles():
    """Trae un diccionario con los equipos en estatus DISPONIBLE (id_estatus = 4)."""
    equipos = {"celulares": [], "laptops": [], "cpus": [], "monitores": [], "tablets": []}
    conn = obtener_conexion()
    if not conn:
        return equipos

    try:
        # 1. Celulares DISPONIBLES (id_estatus_celular = 4)
        df_cel = pd.read_sql("""
            SELECT ic.imei, ic.numero, m.marca_modelo
            FROM inventario_celulares ic
            LEFT JOIN modelos_celulares m ON ic.id_modelo = m.id_modelo
            WHERE ic.id_estatus_celular = 4
        """, conn)
        for _, r in df_cel.iterrows():
            lbl = f"IMEI: {r['imei']} - {r['marca_modelo']} (Línea: {r['numero'] or 'S/N'})"
            equipos["celulares"].append({"id": r['imei'], "label": lbl, "num": r['numero'], "mod": r['marca_modelo']})

        # 2. Laptops DISPONIBLES (id_estatus_laptops = 4)
        df_lap = pd.read_sql("""
            SELECT numero_serie, hostname, marca, modelo, procesador, memoria_ram
            FROM inventario_laptops
            WHERE id_estatus_laptops = 4
        """, conn)
        for _, r in df_lap.iterrows():
            lbl = f"Serie: {r['numero_serie']} - {r['marca']} {r['modelo']} [{r['hostname']}] ({r['procesador']} / {r['memoria_ram']})"
            equipos["laptops"].append({"id": r['numero_serie'], "label": lbl, "host": r['hostname'], "mod": f"{r['marca']} {r['modelo']}"})

        # 3. CPUs DISPONIBLES (id_estatus_cpu = 4)
        df_cpu = pd.read_sql("""
            SELECT hostname, numero_serie, marca, modelo, procesador, memoria_ram
            FROM inventario_cpu
            WHERE id_estatus_cpu = 4
        """, conn)
        for _, r in df_cpu.iterrows():
            lbl = f"Host: {r['hostname']} - Serie: {r['numero_serie']} ({r['marca']} {r['modelo']})"
            equipos["cpus"].append({"id": r['hostname'], "label": lbl, "serie": r['numero_serie'], "mod": f"{r['marca']} {r['modelo']}"})

        # 4. Monitores DISPONIBLES (id_estatus_monitor = 4)
        df_mon = pd.read_sql("""
            SELECT numero_serie, marca, modelo, resolucion
            FROM inventario_monitores
            WHERE id_estatus_monitor = 4
        """, conn)
        for _, r in df_mon.iterrows():
            lbl = f"Serie: {r['numero_serie']} - {r['marca']} {r['modelo']} ({r['resolucion'] or 'Estándar'})"
            equipos["monitores"].append({"id": r['numero_serie'], "label": lbl, "mod": f"{r['marca']} {r['modelo']}"})

        # 5. Tablets DISPONIBLES (id_estatus_tablet = 4)
        df_tab = pd.read_sql("""
            SELECT numero_serie, imei, marca, modelo
            FROM inventario_tablets
            WHERE id_estatus_tablet = 4
        """, conn)
        for _, r in df_tab.iterrows():
            lbl = f"Serie: {r['numero_serie']} - {r['marca']} {r['modelo']} (IMEI: {r['imei'] or 'N/A'})"
            equipos["tablets"].append({"id": r['numero_serie'], "label": lbl, "mod": f"{r['marca']} {r['modelo']}"})

    except Exception as e:
        st.error(f"⚠️ Error al consultar inventario disponible: {e}")
    finally:
        conn.close()

    return equipos

def procesar_asignacion_responsiva(codigo_empleado, cel_sel, lap_sel, cpu_sel, mon_sel, tab_sel):
    """Cambia el estatus a ASIGNADO (id_estatus = 3) y registra las responsivas usando el código exacto de BDD."""
    try:
        conn = obtener_conexion()
        cursor = conn.cursor()
        f_hoy = datetime.now().strftime('%Y-%m-%d')
        
        # Mantenemos el código exacto de empleados para no violar la Foreign Key
        codigo_emp_exacto = str(codigo_empleado).strip()

        # 1. Asignar Celular (Pasa a id_estatus_celular = 3)
        if cel_sel:
            imei = cel_sel['id']
            num = cel_sel['num']
            cursor.execute("UPDATE inventario_celulares SET id_estatus_celular = 3, codigo_empleado = %s WHERE imei = %s", (codigo_emp_exacto, imei))
            cursor.execute("INSERT INTO responsivas_celulares (fecha_entrega, codigo_empleado, numero, imei, id_status) VALUES (%s, %s, %s, %s, 1)", (f_hoy, codigo_emp_exacto, num, imei))

        # 2. Asignar Laptop (Pasa a id_estatus_laptops = 3)
        if lap_sel:
            serie = lap_sel['id']
            cursor.execute("UPDATE inventario_laptops SET id_estatus_laptops = 3, codigo_empleado = %s WHERE numero_serie = %s", (codigo_emp_exacto, serie))
            cursor.execute("INSERT INTO responsivas_laptops (fecha_entrega, codigo_empleado, numero_serie, id_status) VALUES (%s, %s, %s, 1)", (f_hoy, codigo_emp_exacto, serie))

        # 3. Asignar CPU (Pasa a id_estatus_cpu = 3)
        if cpu_sel:
            host = cpu_sel['id']
            cursor.execute("UPDATE inventario_cpu SET id_estatus_cpu = 3, codigo_empleado = %s WHERE hostname = %s", (codigo_emp_exacto, host))
            cursor.execute("INSERT INTO responsivas_cpu (fecha_entrega, codigo_empleado, hostname, id_status) VALUES (%s, %s, %s, 1)", (f_hoy, codigo_emp_exacto, host))

        # 4. Asignar Monitor (Pasa a id_estatus_monitor = 3)
        if mon_sel:
            serie = mon_sel['id']
            cursor.execute("UPDATE inventario_monitores SET id_estatus_monitor = 3, codigo_empleado = %s WHERE numero_serie = %s", (codigo_emp_exacto, serie))
            cursor.execute("INSERT INTO responsivas_monitores (fecha_entrega, codigo_empleado, numero_serie, id_status) VALUES (%s, %s, %s, 1)", (f_hoy, codigo_emp_exacto, serie))

        # 5. Asignar Tablet (Pasa a id_estatus_tablet = 3)
        if tab_sel:
            serie = tab_sel['id']
            cursor.execute("UPDATE inventario_tablets SET id_estatus_tablet = 3, codigo_empleado = %s WHERE numero_serie = %s", (codigo_emp_exacto, serie))
            cursor.execute("INSERT INTO responsivas_tablets (fecha_entrega, codigo_empleado, numero_serie, id_status) VALUES (%s, %s, %s, 1)", (f_hoy, codigo_emp_exacto, serie))

        conn.commit()
        conn.close()
        return True
    except Exception as e:
        st.error(f"⚠️ Error al procesar asignación en BDD: {e}")
        return False

def generar_docx_responsiva(emp_row, cel_sel, lap_sel, cpu_sel, mon_sel, tab_sel):
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = title.add_run("AGROCISA - RESGUARDO Y CARTA RESPONSIVA DE EQUIPO DE CÓMPUTO Y COMUNICACIÓN")
    run_title.bold = True
    run_title.font.size = Pt(13)
    run_title.font.color.rgb = RGBColor(0, 51, 102)

    f_texto = datetime.now().strftime('%d de %B de %Y')
    p_fecha = doc.add_paragraph()
    p_fecha.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_fecha.add_run(f"La Barca, Jalisco a {f_texto}").italic = True

    doc.add_heading("1. DATOS DEL COLABORADOR", level=2)
    p_emp = doc.add_paragraph()
    p_emp.add_run("Código / ID: ").bold = True
    p_emp.add_run(f"{emp_row['codigo_str']}\n")
    p_emp.add_run("Nombre Completo: ").bold = True
    p_emp.add_run(f"{emp_row['nombre_completo']}\n")
    p_emp.add_run("Sucursal / Sede: ").bold = True
    p_emp.add_run(f"{emp_row['sucursal']}\n")
    p_emp.add_run("Departamento / Puesto: ").bold = True
    p_emp.add_run(f"{emp_row['departamento']} - {emp_row['puesto']}")

    doc.add_heading("2. DETALLE DE HARDWARE ASIGNADO", level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Tipo de Hardware'
    hdr_cells[1].text = 'Identificador / Serie / IMEI'
    hdr_cells[2].text = 'Descripción / Modelo'
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].font.bold = True

    items = []
    if cel_sel: items.append(("Celular", cel_sel['id'], f"{cel_sel['mod']} (Línea: {cel_sel['num'] or 'S/N'})"))
    if lap_sel: items.append(("Laptop", lap_sel['id'], f"{lap_sel['mod']} [Host: {lap_sel['host']}]"))
    if cpu_sel: items.append(("CPU Desktop", cpu_sel['serie'], f"{cpu_sel['mod']} [Host: {cpu_sel['id']}]"))
    if mon_sel: items.append(("Monitor", mon_sel['id'], mon_sel['mod']))
    if tab_sel: items.append(("Tablet", tab_sel['id'], tab_sel['mod']))

    for tipo, ident, desc in items:
        row_cells = table.add_row().cells
        row_cells[0].text = tipo
        row_cells[1].text = str(ident)
        row_cells[2].text = str(desc)

    doc.add_heading("3. COMPROMISOS Y TÉRMINOS DE USO", level=2)
    p_cla = doc.add_paragraph()
    p_cla.add_run("1. El colaborador reconoce haber recibido el equipo descrito en perfectas condiciones operativas y físicas para el desempeño exclusivo de sus funciones laborales.\n")
    p_cla.add_run("2. Es responsabilidad del usuario la custodia, cuidado y buen uso del hardware y software instalado.\n")
    p_cla.add_run("3. En caso de extravío, daño por negligencia o robo, el usuario deberá notificar de inmediato al Departamento de TI.")

    doc.add_paragraph("\n\n")
    p_firmas = doc.add_paragraph()
    p_firmas.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_firmas.add_run("_________________________________________\n").bold = True
    p_firmas.add_run(f"{emp_row['nombre_completo']}\n").bold = True
    p_firmas.add_run("FIRMA DE CONFORMIDAD DEL COLABORADOR").font.size = Pt(9)

    target = io.BytesIO()
    doc.save(target)
    target.seek(0)
    return target

def render():
    aplicar_estilos_pantalla()
    st.subheader("📄 Generación de Responsivas y Asignación de Hardware")

    df_emp = obtener_empleados_activos_df()
    dict_equipos = obtener_equipos_disponibles()

    if df_emp.empty:
        st.warning("⚠️ No se encontraron empleados activos en la base de datos.")
        return

    col1, col2 = st.columns([2, 1])
    with col1:
        lista_emp = [f"{r['codigo_str']} - {r['nombre_completo']} ({r['sucursal']})" for _, r in df_emp.iterrows()]
        emp_sel_str = st.selectbox("Selecciona el colaborador a asignar:", lista_emp)
        codigo_sel = emp_sel_str.split(" - ")[0]
        emp_row = df_emp[df_emp["codigo_str"] == codigo_sel].iloc[0]

    with col2:
        st.info(f"**Sucursal:** {emp_row['sucursal']}\n\n**Depto:** {emp_row['departamento']}\n\n**Puesto:** {emp_row['puesto']}")

    st.divider()
    st.markdown("### 📦 Selecciona el Hardware DISPONIBLE a Asignar")

    c1, c2 = st.columns(2)
    with c1:
        opts_cel = ["-- Ninguno --"] + [x["label"] for x in dict_equipos["celulares"]]
        sel_cel_txt = st.selectbox("📱 Celulares Disponibles:", opts_cel)
        obj_cel = next((x for x in dict_equipos["celulares"] if x["label"] == sel_cel_txt), None)

        opts_lap = ["-- Ninguno --"] + [x["label"] for x in dict_equipos["laptops"]]
        sel_lap_txt = st.selectbox("💻 Laptops Disponibles:", opts_lap)
        obj_lap = next((x for x in dict_equipos["laptops"] if x["label"] == sel_lap_txt), None)

        opts_tab = ["-- Ninguno --"] + [x["label"] for x in dict_equipos["tablets"]]
        sel_tab_txt = st.selectbox("📱 Tablets Disponibles:", opts_tab)
        obj_tab = next((x for x in dict_equipos["tablets"] if x["label"] == sel_tab_txt), None)

    with c2:
        opts_cpu = ["-- Ninguno --"] + [x["label"] for x in dict_equipos["cpus"]]
        sel_cpu_txt = st.selectbox("🖥️ CPUs Disponibles:", opts_cpu)
        obj_cpu = next((x for x in dict_equipos["cpus"] if x["label"] == sel_cpu_txt), None)

        opts_mon = ["-- Ninguno --"] + [x["label"] for x in dict_equipos["monitores"]]
        sel_mon_txt = st.selectbox("🖥️ Monitores Disponibles:", opts_mon)
        obj_mon = next((x for x in dict_equipos["monitores"] if x["label"] == sel_mon_txt), None)

    st.divider()

    if not (obj_cel or obj_lap or obj_cpu or obj_mon or obj_tab):
        st.warning("👉 Selecciona al menos un equipo de las listas para habilitar la asignación.")
    else:
        if st.button("🚀 Confirmar Asignación en BDD y Generar Documento DOCX", type="primary"):
            if procesar_asignacion_responsiva(codigo_sel, obj_cel, obj_lap, obj_cpu, obj_mon, obj_tab):
                st.session_state["docx_generado"] = generar_docx_responsiva(emp_row, obj_cel, obj_lap, obj_cpu, obj_mon, obj_tab)
                st.session_state["nombre_docx"] = f"Responsiva_{codigo_sel}_{emp_row['nombre_completo'].replace(' ', '_')}.docx"
                st.toast("¡Asignación guardada en BDD!", icon="🎉")
                st.rerun()

    if "docx_generado" in st.session_state:
        st.success("✅ ¡Transacción completada! El hardware cambió a ASIGNADO y el documento está listo:")
        st.download_button(
            label="📥 Descargar Carta Responsiva (.docx)",
            data=st.session_state["docx_generado"],
            file_name=st.session_state["nombre_docx"],
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            type="primary"
        )
        if st.button("🔄 Nueva Asignación"):
            del st.session_state["docx_generado"]
            del st.session_state["nombre_docx"]
            st.rerun()